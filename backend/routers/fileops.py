"""
File operations router:
  POST /fileops/encrypt-text          — encrypt a text block
  POST /fileops/decrypt-text          — decrypt an encrypted token
  POST /fileops/apply-encryption/{id} — replace text in doc with encrypted token
  POST /fileops/apply-decryption/{id} — replace token in doc with plaintext
  GET  /fileops/content/{id}          — get document text for viewer
  POST /fileops/lock/{id}             — password-lock a document
  POST /fileops/unlock/{id}           — unlock a document
  POST /fileops/convert/{id}          — convert document format
"""
import io
import uuid
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from pydantic import BaseModel

from models.database import get_db
from models.models import Document, User
from routers.deps import get_current_user
from utils.crypto import encrypt_file, decrypt_file
from utils.file_ops import (
    encrypt_text_block, decrypt_text_block,
    lock_file, unlock_file,
    pdf_to_docx, docx_to_pdf, docx_to_txt, txt_to_pdf,
)
router = APIRouter()
ENCRYPTED_FILES_DIR = Path(__file__).parent.parent / "encrypted_files"


# ── Schemas ───────────────────────────────────────────────────────────────────

class TextEncryptRequest(BaseModel):
    text: str

class TextDecryptRequest(BaseModel):
    token: str

class ApplyEncryptionRequest(BaseModel):
    original_text: str
    encrypted_token: str

class LockRequest(BaseModel):
    password: str

class ConvertRequest(BaseModel):
    target_format: str


# ── Standalone text encrypt/decrypt ──────────────────────────────────────────

@router.post("/encrypt-text")
async def encrypt_text(body: TextEncryptRequest, current_user: User = Depends(get_current_user)):
    token = encrypt_text_block(body.text, current_user.id)
    return {"token": token, "original_length": len(body.text)}


@router.post("/decrypt-text")
async def decrypt_text(body: TextDecryptRequest, current_user: User = Depends(get_current_user)):
    try:
        plaintext = decrypt_text_block(body.token, current_user.id)
        return {"text": plaintext}
    except Exception:
        raise HTTPException(status_code=400, detail="Decryption failed — invalid token or wrong user")


# ── Apply encryption/decryption to stored document ───────────────────────────

async def _modify_document(doc_id: int, original_text: str, replacement: str,
                            db: AsyncSession, current_user: User) -> dict:
    """Find original_text in document and replace with replacement. Saves back."""
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.status == "locked":
        raise HTTPException(status_code=400, detail="Unlock the document before editing")

    mime = doc.mime_type or ""
    file_bytes = decrypt_file(doc.enc_path, current_user.id)

    if mime == "application/pdf":
        raise HTTPException(
            status_code=400,
            detail="PDF files cannot be modified. Convert to DOCX or TXT first."
        )
    elif mime == "text/plain":
        text = file_bytes.decode("utf-8", errors="replace")
        if original_text not in text:
            raise HTTPException(status_code=400,
                detail=f"Text not found in document (case-sensitive).")
        new_bytes = text.replace(original_text, replacement, 1).encode("utf-8")

    elif "wordprocessingml" in mime:
        import docx as _docx, io as _io
        document = _docx.Document(_io.BytesIO(file_bytes))
        full_text = "\n".join(p.text for p in document.paragraphs)
        if original_text not in full_text:
            raise HTTPException(status_code=400,
                detail=f"Text not found in document (case-sensitive).")
        replaced = False
        for para in document.paragraphs:
            if original_text in para.text and not replaced:
                new_para_text = para.text.replace(original_text, replacement, 1)
                for i, run in enumerate(para.runs):
                    run.text = new_para_text if i == 0 else ""
                replaced = True
                break
        buf = _io.BytesIO()
        document.save(buf)
        new_bytes = buf.getvalue()
    else:
        raise HTTPException(status_code=400, detail="Only TXT and DOCX supported")

    old_path = Path(doc.enc_path)
    new_enc_path = encrypt_file(new_bytes, current_user.id, f"{uuid.uuid4().hex}_{doc.filename}")
    doc.enc_path = str(new_enc_path)
    await db.commit()
    if old_path.exists():
        old_path.unlink()
    return {"message": "Document updated successfully"}


@router.post("/apply-encryption/{doc_id}")
async def apply_text_encryption(
    doc_id: int, body: ApplyEncryptionRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return await _modify_document(doc_id, body.original_text, body.encrypted_token, db, current_user)


@router.post("/apply-decryption/{doc_id}")
async def apply_text_decryption(
    doc_id: int, body: ApplyEncryptionRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # original_text = the [[ENC:...]] token, encrypted_token = the plaintext to restore
    return await _modify_document(doc_id, body.original_text, body.encrypted_token, db, current_user)


# ── Document content viewer ───────────────────────────────────────────────────

@router.get("/content/{doc_id}")
async def get_document_content(
    doc_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.status == "locked":
        raise HTTPException(status_code=400, detail="Document is locked")

    mime = doc.mime_type or ""
    if mime == "application/pdf":
        raise HTTPException(status_code=400, detail="PDF preview not supported — convert to DOCX or TXT first")

    file_bytes = decrypt_file(doc.enc_path, current_user.id)

    if mime == "text/plain":
        return {"content": file_bytes.decode("utf-8", errors="replace")}
    elif "wordprocessingml" in mime:
        import docx as _docx, io as _io
        document = _docx.Document(_io.BytesIO(file_bytes))
        return {"content": "\n".join(p.text for p in document.paragraphs)}
    else:
        raise HTTPException(status_code=400, detail="Preview only available for TXT and DOCX")


# ── File locking ──────────────────────────────────────────────────────────────

@router.post("/lock/{doc_id}")
async def lock_document(
    doc_id: int, body: LockRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.status == "locked":
        raise HTTPException(status_code=400, detail="Document is already locked")

    file_bytes = decrypt_file(doc.enc_path, current_user.id)
    locked_bytes = lock_file(file_bytes, body.password)

    ENCRYPTED_FILES_DIR.mkdir(parents=True, exist_ok=True)
    locked_path = ENCRYPTED_FILES_DIR / f"locked_{uuid.uuid4().hex}.lck"
    locked_path.write_bytes(locked_bytes)

    old_path = Path(doc.enc_path)
    doc.enc_path = str(locked_path)
    doc.status = "locked"
    await db.commit()
    if old_path.exists():
        old_path.unlink()

    return {"message": "Document locked successfully", "status": "locked"}


@router.post("/unlock/{doc_id}")
async def unlock_document(
    doc_id: int, body: LockRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.status != "locked":
        raise HTTPException(status_code=400, detail="Document is not locked")

    locked_bytes = Path(doc.enc_path).read_bytes()
    try:
        file_bytes = unlock_file(locked_bytes, body.password)
    except ValueError:
        raise HTTPException(status_code=400, detail="Incorrect password")

    new_enc_path = encrypt_file(file_bytes, current_user.id, f"{uuid.uuid4().hex}_{doc.filename}")
    old_path = Path(doc.enc_path)
    doc.enc_path = str(new_enc_path)
    doc.status = "ready"
    await db.commit()
    if old_path.exists():
        old_path.unlink()

    return {"message": "Document unlocked successfully", "status": "ready"}


# ── File conversion ───────────────────────────────────────────────────────────

CONVERSION_MAP = {
    ("application/pdf", "docx"): (pdf_to_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx"),
    ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "pdf"): (docx_to_pdf, "application/pdf", ".pdf"),
    ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "txt"): (docx_to_txt, "text/plain", ".txt"),
    ("text/plain", "pdf"): (txt_to_pdf, "application/pdf", ".pdf"),
}


@router.post("/convert/{doc_id}")
async def convert_document(
    doc_id: int, body: ConvertRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.status == "locked":
        raise HTTPException(status_code=400, detail="Unlock the document before converting")

    key = (doc.mime_type, body.target_format.lower())
    if key not in CONVERSION_MAP:
        raise HTTPException(status_code=400, detail=f"Unsupported conversion for {doc.mime_type} → {body.target_format}")

    converter_fn, out_mime, out_ext = CONVERSION_MAP[key]
    file_bytes = decrypt_file(doc.enc_path, current_user.id)

    try:
        converted = converter_fn(file_bytes)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {e}")

    filename = f"{Path(doc.original_name).stem}{out_ext}"
    return StreamingResponse(
        io.BytesIO(converted),
        media_type=out_mime,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── PDF password protection ───────────────────────────────────────────────────

class PdfPasswordRequest(BaseModel):
    password: str


@router.post("/download-protected/{doc_id}")
async def download_password_protected(
    doc_id: int,
    body: PdfPasswordRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Download a PDF with password protection (like Aadhaar PDF).
    Non-PDF files are wrapped in a password-protected ZIP instead.
    """
    result = await db.execute(
        select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    if doc.status == "locked":
        raise HTTPException(status_code=400, detail="Unlock the document first")

    file_bytes = decrypt_file(doc.enc_path, current_user.id)
    mime = doc.mime_type or ""

    if mime == "application/pdf":
        # Apply PDF owner/user password using PyPDF2
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            writer = PyPDF2.PdfWriter()
            for page in reader.pages:
                writer.add_page(page)
            writer.encrypt(body.password, algorithm="AES-256")
            buf = io.BytesIO()
            writer.write(buf)
            protected_bytes = buf.getvalue()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF protection failed: {e}")

        return StreamingResponse(
            io.BytesIO(protected_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{doc.original_name}"'},
        )
    else:
        # Non-PDF: wrap in password-protected ZIP
        import pyzipper
        buf = io.BytesIO()
        with pyzipper.AESZipFile(buf, "w", compression=pyzipper.ZIP_DEFLATED,
                                  encryption=pyzipper.WZ_AES) as zf:
            zf.setpassword(body.password.encode())
            zf.writestr(doc.original_name, file_bytes)
        zip_name = f"{Path(doc.original_name).stem}_protected.zip"
        return StreamingResponse(
            io.BytesIO(buf.getvalue()),
            media_type="application/zip",
            headers={"Content-Disposition": f'attachment; filename="{zip_name}"'},
        )


# ── ZIP download ──────────────────────────────────────────────────────────────

class ZipRequest(BaseModel):
    doc_ids: list[int]
    password: str = ""  # optional password protection


@router.post("/download-zip")
async def download_as_zip(
    body: ZipRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Package one or more documents into a ZIP file (optionally password-protected)."""
    if not body.doc_ids:
        raise HTTPException(status_code=400, detail="No documents selected")

    buf = io.BytesIO()
    use_password = bool(body.password.strip())

    if use_password:
        import pyzipper
        zf_ctx = pyzipper.AESZipFile(buf, "w", compression=pyzipper.ZIP_DEFLATED,
                                      encryption=pyzipper.WZ_AES)
    else:
        import zipfile
        zf_ctx = zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED)

    added = 0
    with zf_ctx as zf:
        if use_password:
            zf.setpassword(body.password.encode())
        for doc_id in body.doc_ids:
            result = await db.execute(
                select(Document).where(Document.id == doc_id, Document.user_id == current_user.id)
            )
            doc = result.scalar_one_or_none()
            if not doc or doc.status == "locked":
                continue
            try:
                file_bytes = decrypt_file(doc.enc_path, current_user.id)
                zf.writestr(doc.original_name, file_bytes)
                added += 1
            except Exception:
                continue

    if added == 0:
        raise HTTPException(status_code=400, detail="No accessible documents found")

    suffix = "_protected" if use_password else ""
    return StreamingResponse(
        io.BytesIO(buf.getvalue()),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="documents{suffix}.zip"'},
    )
