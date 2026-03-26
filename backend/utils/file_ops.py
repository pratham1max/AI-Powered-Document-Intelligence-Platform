"""
File operations utility:
- Selective text encryption/decryption (AES-256-GCM)
- File locking with user password
- File format conversion (PDF↔DOCX, DOCX→TXT)
"""
import os
import io
import base64
import hashlib
from cryptography.hazmat.primitives.ciphers.aead import AESGCM

MASTER_SECRET = os.getenv("MASTER_SECRET", "").encode()
LOCK_MARKER_START = "[[ENC:"
LOCK_MARKER_END = "]]"


# ── Selective text encryption ─────────────────────────────────────────────────

def _derive_text_key(user_id: int) -> bytes:
    return hashlib.sha256(MASTER_SECRET + f":textenc:{user_id}".encode()).digest()


def encrypt_text_block(plaintext: str, user_id: int) -> str:
    """Encrypt a text block → returns a compact base64 token wrapped in markers."""
    key = _derive_text_key(user_id)
    nonce = os.urandom(12)
    ct = AESGCM(key).encrypt(nonce, plaintext.encode(), None)
    token = base64.urlsafe_b64encode(nonce + ct).decode()
    return f"{LOCK_MARKER_START}{token}{LOCK_MARKER_END}"


def decrypt_text_block(token_with_markers: str, user_id: int) -> str:
    """Decrypt a marked encrypted token back to plaintext."""
    if not (token_with_markers.startswith(LOCK_MARKER_START) and
            token_with_markers.endswith(LOCK_MARKER_END)):
        raise ValueError("Not a valid encrypted block")
    token = token_with_markers[len(LOCK_MARKER_START):-len(LOCK_MARKER_END)]
    raw = base64.urlsafe_b64decode(token)
    nonce, ct = raw[:12], raw[12:]
    key = _derive_text_key(user_id)
    return AESGCM(key).decrypt(nonce, ct, None).decode()


# ── File locking (password-based) ─────────────────────────────────────────────

def _derive_lock_key(password: str, salt: bytes) -> bytes:
    from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
    from cryptography.hazmat.primitives import hashes
    kdf = PBKDF2HMAC(algorithm=hashes.SHA256(), length=32, salt=salt, iterations=100_000)
    return kdf.derive(password.encode())


def lock_file(file_bytes: bytes, password: str) -> bytes:
    """Encrypt file bytes with a user-supplied password. Returns salt+nonce+ct."""
    salt = os.urandom(16)
    nonce = os.urandom(12)
    key = _derive_lock_key(password, salt)
    ct = AESGCM(key).encrypt(nonce, file_bytes, None)
    return salt + nonce + ct


def unlock_file(locked_bytes: bytes, password: str) -> bytes:
    """Decrypt a password-locked file. Raises ValueError on wrong password."""
    if len(locked_bytes) < 28:
        raise ValueError("Invalid locked file")
    salt, nonce, ct = locked_bytes[:16], locked_bytes[16:28], locked_bytes[28:]
    key = _derive_lock_key(password, salt)
    try:
        return AESGCM(key).decrypt(nonce, ct, None)
    except Exception:
        raise ValueError("Incorrect password or corrupted file")


# ── File conversion ───────────────────────────────────────────────────────────

def pdf_to_docx(pdf_bytes: bytes) -> bytes:
    """Convert PDF bytes to DOCX bytes."""
    import tempfile, os
    from pdf2docx import Converter
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(pdf_bytes)
        pdf_path = f.name
    docx_path = pdf_path.replace(".pdf", ".docx")
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        with open(docx_path, "rb") as f:
            return f.read()
    finally:
        os.unlink(pdf_path)
        if os.path.exists(docx_path):
            os.unlink(docx_path)


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF using reportlab (text extraction → PDF)."""
    import docx
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import cm

    doc = docx.Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []
    for text in paragraphs:
        story.append(Paragraph(text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["Normal"]))
        story.append(Spacer(1, 6))
    pdf.build(story)
    return buf.getvalue()


def docx_to_txt(docx_bytes: bytes) -> bytes:
    """Extract plain text from DOCX."""
    import docx
    doc = docx.Document(io.BytesIO(docx_bytes))
    text = "\n".join(p.text for p in doc.paragraphs)
    return text.encode("utf-8")


def txt_to_pdf(txt_bytes: bytes) -> bytes:
    """Convert plain text to PDF."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import cm

    text = txt_bytes.decode("utf-8", errors="replace")
    buf = io.BytesIO()
    pdf = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []
    for line in text.splitlines():
        safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe or "&nbsp;", styles["Normal"]))
        story.append(Spacer(1, 4))
    pdf.build(story)
    return buf.getvalue()
